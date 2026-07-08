import os
from pathlib import Path
import pypdf
from docx import Document

def extract_candidate_name(filepath: Path) -> str:
    """
    Extracts the candidate's canonical name from the resume filename.
    Format: '<Description> - <Candidate Name>.<ext>'
    Example: 'ASHOK_Reddy_RESUME - M Ashok reddy.pdf' -> 'M Ashok reddy'
    """
    filename = filepath.stem  # Strip extension
    if " - " in filename:
        parts = filename.split(" - ")
        # Take the part after the hyphen
        name = parts[-1].strip()
    else:
        name = filename.strip()
    return name

def parse_pdf(filepath: Path) -> str:
    """Parses a PDF file and returns its text content."""
    text_content = []
    try:
        reader = pypdf.PdfReader(filepath)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    except Exception as e:
        print(f"Error parsing PDF file {filepath}: {e}")
        raise e
    
    return "\n\n".join(text_content)

def parse_docx(filepath: Path) -> str:
    """Parses a DOCX file and returns its text content."""
    text_content = []
    try:
        doc = Document(filepath)
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        # Read tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
    except Exception as e:
        print(f"Error parsing DOCX file {filepath}: {e}")
        raise e
    
    return "\n\n".join(text_content)

def parse_resume(filepath: Path) -> tuple[str, str]:
    """
    Parses a resume (PDF or DOCX) and returns (candidate_name, text_content).
    """
    ext = filepath.suffix.lower()
    candidate_name = extract_candidate_name(filepath)
    
    if ext == ".pdf":
        text_content = parse_pdf(filepath)
    elif ext == ".docx":
        text_content = parse_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    return candidate_name, text_content
